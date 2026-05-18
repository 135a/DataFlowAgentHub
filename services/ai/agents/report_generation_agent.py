import logging
import os
import re
import uuid
import hashlib
import pandas as pd
from datetime import datetime
from orchestrator.state import AgentState
from orchestrator.tracing import report_run_step

# PDF & DOCX generation
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# 报告输出目录（环境变量化）
REPORT_OUTPUT_DIR = os.getenv("REPORT_OUTPUT_DIR", "/tmp/reports")


def report_generation_node(state: AgentState) -> dict:
    logger.info("Executing Report Generation Agent")
    run_id = state.get("run_id", "")
    report_run_step(run_id, "report_generation_agent", "running", "Generating report...")

    # ----------------------------------------------------------------
    # 路径遍历防护：确保 run_id 不包含路径分隔符
    # ----------------------------------------------------------------
    raw_run_id = state.get("run_id") or str(uuid.uuid4())
    if not re.match(r'^[a-fA-F0-9-]{36}$', raw_run_id):
        safe_run_id = hashlib.sha256(raw_run_id.encode()).hexdigest()[:16]
    else:
        safe_run_id = raw_run_id

    # 生成 Markdown 报告
    md_lines = [
        "# Data Analysis Report",
        f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"\n## Request\n{state.get('user_input', 'N/A')}",
        f"\n## Analysis Summary\n{state.get('analysis_summary', 'No summary available.')}",
        f"\n## Data Extract",
    ]

    result_data = state.get("nl2sql_result", [])
    if result_data:
        df = pd.DataFrame(result_data)
        md_lines.append(df.head(10).to_markdown(index=False))
        if len(result_data) > 10:
            md_lines.append(f"\n*(Showing top 10 rows of {len(result_data)} total)*")
    else:
        md_lines.append("*No data returned.*")

    # 如果存在图表路径，添加图表可视化章节
    chart_paths = state.get("chart_paths", [])
    if chart_paths:
        md_lines.append("\n## 数据可视化")
        for p in chart_paths:
            filename = os.path.basename(p)
            md_lines.append(f"\n![chart](./{filename})")
        md_lines.append(f"\n*共 {len(chart_paths)} 个图表*")

    final_md = "\n".join(md_lines)

    # 保存到 report 目录（环境变量化）
    os.makedirs(REPORT_OUTPUT_DIR, exist_ok=True)

    # 生成 Markdown 报告
    md_path = os.path.join(REPORT_OUTPUT_DIR, f"{safe_run_id}.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(final_md)
    logger.info("Markdown report saved to %s", md_path)

    # 生成 PDF 报告
    if result_data:
        try:
            pdf_path = os.path.join(REPORT_OUTPUT_DIR, f"{safe_run_id}.pdf")
            _generate_pdf(final_md, df, pdf_path)
            logger.info("PDF report saved to %s", pdf_path)
        except Exception as e:
            logger.error("Failed to generate PDF report: %s", e)

    # 生成 DOCX 报告
    if result_data:
        try:
            docx_path = os.path.join(REPORT_OUTPUT_DIR, f"{safe_run_id}.docx")
            _generate_docx(final_md, df, docx_path)
            logger.info("DOCX report saved to %s", docx_path)
        except Exception as e:
            logger.error("Failed to generate DOCX report: %s", e)

    report_run_step(run_id, "report_generation_agent", "succeeded", "Report generated", f"Files saved for run_id: {safe_run_id}")
    return {"final_report": final_md}


def _generate_pdf(md_content: str, df: pd.DataFrame, output_path: str):
    """使用 fpdf2 生成 PDF 报告"""
    pdf = FPDF()
    pdf.add_page()

    # 标题
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Data Analysis Report", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    # 时间戳
    pdf.set_font("Helvetica", "", 8)
    pdf.cell(0, 6, f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", new_x="LMARGIN", new_y="NEXT")

    # 请求摘要
    for section_title, prefix in [("Request", "## Request"), ("Analysis Summary", "## Analysis Summary")]:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, section_title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        # 从 markdown 中提取段落内容
        for line in md_content.split("\n"):
            if line.startswith(prefix):
                continue
            if line.startswith("#") or line.startswith("|") or line.startswith("*"):
                continue
            if line.strip():
                pdf.multi_cell(0, 5, line.strip())
                pdf.ln(2)

    # 数据表格
    if not df.empty:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Data Extract", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        # 表格列
        pdf.set_font("Helvetica", "B", 8)
        col_count = len(df.columns)
        col_width = min(180 / max(col_count, 1), 60)
        for col in df.columns:
            pdf.cell(col_width, 6, str(col)[:20], border=1, align="C")
        pdf.ln()

        # 表格行（最多显示 10 行）
        pdf.set_font("Helvetica", "", 7)
        for _, row in df.head(10).iterrows():
            for val in row:
                text = str(val)[:30] if val is not None else ""
                pdf.cell(col_width, 5, text, border=1)
            pdf.ln()

        if len(df) > 10:
            pdf.set_font("Helvetica", "I", 8)
            pdf.cell(0, 6, f"(Showing top 10 rows of {len(df)} total)", new_x="LMARGIN", new_y="NEXT")

    pdf.output(output_path)


def _generate_docx(md_content: str, df: pd.DataFrame, output_path: str):
    """使用 python-docx 生成 DOCX 报告"""
    doc = Document()

    # 标题
    title = doc.add_heading("Data Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 时间戳
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)

    # 请求
    doc.add_heading("Request", level=1)
    for line in md_content.split("\n"):
        if line.startswith("## Request"):
            continue
        if line.startswith("#"):
            break
        if line.strip() and not line.startswith("|") and not line.startswith("*"):
            doc.add_paragraph(line.strip())

    # 分析摘要
    doc.add_heading("Analysis Summary", level=1)
    in_summary = False
    for line in md_content.split("\n"):
        if line.startswith("## Analysis Summary"):
            in_summary = True
            continue
        if in_summary and (line.startswith("## ") and "Analysis Summary" not in line):
            break
        if in_summary and line.strip() and not line.startswith("|"):
            doc.add_paragraph(line.strip())

    # 数据表格
    if not df.empty:
        doc.add_heading("Data Extract", level=1)
        table = doc.add_table(rows=min(len(df), 10) + 1, cols=len(df.columns))
        table.style = "Light Grid Accent 1"

        # 表头
        for i, col in enumerate(df.columns):
            cell = table.rows[0].cells[i]
            cell.text = str(col)

        # 数据行
        for row_idx, (_, row) in enumerate(df.head(10).iterrows()):
            for col_idx, val in enumerate(row):
                table.rows[row_idx + 1].cells[col_idx].text = str(val) if val is not None else ""

        if len(df) > 10:
            doc.add_paragraph(f"(Showing top 10 rows of {len(df)} total)")

    doc.save(output_path)
